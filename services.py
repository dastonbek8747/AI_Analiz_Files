import os
import platform
from fpdf import FPDF
import pandas as pd
from PIL import Image
import json
import PyPDF2
import docx
import time

try:
    import pythoncom
    from win32com import client

    WINDOWS_AVAILABLE = True
except ImportError:
    WINDOWS_AVAILABLE = False


class UniversalConverter:
    """Universal fayl konvertor klassi"""

    def __init__(self):
        self.base_path = "Files"
        self.output_path = "Output"
        os.makedirs(self.base_path, exist_ok=True)
        os.makedirs(self.output_path, exist_ok=True)

    def get_file_path(self, filename: str, is_output: bool = False) -> str:
        """Fayl yo'lini olish"""
        folder = self.output_path if is_output else self.base_path
        return os.path.abspath(os.path.join(folder, filename))

    def create_output_filename(self, input_file: str, output_ext: str) -> str:
        """Output fayl nomini yaratish"""
        base_name = os.path.splitext(os.path.basename(input_file))[0]
        # 'saved_' prefiksini olib tashlash
        if base_name.startswith('saved_'):
            base_name = base_name[6:]
        return os.path.join(self.output_path, f"{base_name}.{output_ext}")

    def safe_read_excel(self, file_path, max_retries=3):
        """Excel faylni xavfsiz o'qish"""
        for attempt in range(max_retries):
            try:
                return pd.read_excel(file_path, engine='openpyxl')
            except PermissionError:
                if attempt < max_retries - 1:
                    time.sleep(0.5)
                    continue
                else:
                    raise
            except Exception as e:
                if attempt < max_retries - 1:
                    time.sleep(0.5)
                    continue
                else:
                    raise e


# ==================== PDF ga konvertatsiya ====================

def to_pdf(filename: str, file_type: str = None, use_com: bool = False) -> str:
    """Har qanday faylni PDF ga aylantirish

    Args:
        filename: Fayl nomi
        file_type: Fayl turi (agar None bo'lsa, avtomatik aniqlanadi)
        use_com: Windows COM API ishlatish (False = pandas/FPDF, True = Microsoft Office)
    """
    converter = UniversalConverter()
    file_path = converter.get_file_path(f"saved_{filename}")

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Fayl topilmadi: {file_path}")

    # Fayl turini aniqlash
    if file_type is None:
        file_type = filename.split('.')[-1].lower()

    # Tegishli konvertor funksiyasini chaqirish
    converters = {
        'docx': _simple_word_to_pdf,
        'doc': _simple_word_to_pdf,
        'xlsx': _simple_excel_to_pdf,
        'xls': _simple_excel_to_pdf,
        'csv': CSV_to_pdf,
        'txt': TXT_to_pdf,
        'json': JSON_to_pdf,
        'jpg': Image_to_pdf,
        'jpeg': Image_to_pdf,
        'png': Image_to_pdf,
        'bmp': Image_to_pdf,
        'gif': Image_to_pdf,
        'webp': Image_to_pdf,
    }

    converter_func = converters.get(file_type)
    if converter_func:
        return converter_func(filename)
    else:
        raise ValueError(f"Qo'llab-quvvatlanmaydigan fayl turi: {file_type}")


def _simple_excel_to_pdf(filename: str) -> str:
    """Excel ni PDF ga oddiy usulda aylantirish"""
    converter = UniversalConverter()
    file_path = converter.get_file_path(f"saved_{filename}")
    output_file = converter.create_output_filename(file_path, "pdf")

    # Excel faylni xavfsiz o'qish
    df = converter.safe_read_excel(file_path)
    time.sleep(0.2)  # Faylni yopilishini kutish

    return _dataframe_to_pdf(df, output_file, filename)


def _simple_word_to_pdf(filename: str) -> str:
    """Word ni PDF ga oddiy usulda aylantirish"""
    converter = UniversalConverter()
    file_path = converter.get_file_path(f"saved_{filename}")
    output_file = converter.create_output_filename(file_path, "pdf")

    max_retries = 3
    for attempt in range(max_retries):
        try:
            doc = docx.Document(file_path)
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font("Arial", size=11)

            for para in doc.paragraphs:
                if para.text.strip():
                    try:
                        pdf.multi_cell(0, 8, para.text)
                    except:
                        cleaned_text = para.text.encode('latin-1', 'ignore').decode('latin-1')
                        pdf.multi_cell(0, 8, cleaned_text)
                    pdf.ln(2)

            pdf.output(output_file)
            return output_file

        except PermissionError:
            if attempt < max_retries - 1:
                time.sleep(0.5)
                continue
            else:
                raise Exception(f"Fayl band: {file_path}")
        except Exception as e:
            if attempt < max_retries - 1:
                time.sleep(0.5)
                continue
            else:
                raise Exception(f"Word faylni o'qishda xatolik: {str(e)}")


def CSV_to_pdf(filename: str) -> str:
    """CSV faylni PDF ga aylantirish"""
    converter = UniversalConverter()
    file_path = converter.get_file_path(f"saved_{filename}")
    output_file = converter.create_output_filename(file_path, "pdf")

    max_retries = 3
    for attempt in range(max_retries):
        try:
            df = pd.read_csv(file_path)
            time.sleep(0.2)
            return _dataframe_to_pdf(df, output_file, filename)
        except PermissionError:
            if attempt < max_retries - 1:
                time.sleep(0.5)
                continue
            else:
                raise
        except Exception as e:
            if attempt < max_retries - 1:
                time.sleep(0.5)
                continue
            else:
                raise e


def TXT_to_pdf(filename: str) -> str:
    """TXT faylni PDF ga aylantirish"""
    converter = UniversalConverter()
    file_path = converter.get_file_path(f"saved_{filename}")
    output_file = converter.create_output_filename(file_path, "pdf")

    max_retries = 3
    for attempt in range(max_retries):
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font("Arial", size=11)

            for line in content.splitlines():
                try:
                    pdf.multi_cell(0, 8, line)
                except:
                    pdf.multi_cell(0, 8, line.encode('latin-1', 'ignore').decode('latin-1'))

            pdf.output(output_file)
            return output_file

        except PermissionError:
            if attempt < max_retries - 1:
                time.sleep(0.5)
                continue
            else:
                raise
        except Exception as e:
            if attempt < max_retries - 1:
                time.sleep(0.5)
                continue
            else:
                raise e


def JSON_to_pdf(filename: str) -> str:
    """JSON faylni PDF ga aylantirish"""
    converter = UniversalConverter()
    file_path = converter.get_file_path(f"saved_{filename}")
    output_file = converter.create_output_filename(file_path, "pdf")

    max_retries = 3
    for attempt in range(max_retries):
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font("Arial", size=10)

            pdf.cell(0, 10, f"JSON Fayl: {filename}", ln=True)
            pdf.ln(5)

            json_str = json.dumps(data, indent=2, ensure_ascii=False)
            for line in json_str.splitlines():
                try:
                    pdf.multi_cell(0, 6, line)
                except:
                    pdf.multi_cell(0, 6, line.encode('latin-1', 'ignore').decode('latin-1'))

            pdf.output(output_file)
            return output_file

        except PermissionError:
            if attempt < max_retries - 1:
                time.sleep(0.5)
                continue
            else:
                raise
        except Exception as e:
            if attempt < max_retries - 1:
                time.sleep(0.5)
                continue
            else:
                raise e


def Image_to_pdf(filename: str) -> str:
    """Rasm faylni PDF ga aylantirish"""
    converter = UniversalConverter()
    file_path = converter.get_file_path(f"saved_{filename}")
    output_file = converter.create_output_filename(file_path, "pdf")

    max_retries = 3
    for attempt in range(max_retries):
        try:
            img = Image.open(file_path)

            if img.mode != 'RGB':
                img = img.convert('RGB')

            temp_path = output_file.replace('.pdf', '_temp.jpg')
            img.save(temp_path, 'JPEG')

            pdf = FPDF()
            pdf.add_page()

            img_width, img_height = img.size
            page_width = 210 - 20
            page_height = 297 - 20

            ratio = min(page_width / img_width, page_height / img_height)
            new_width = img_width * ratio
            new_height = img_height * ratio

            x = (210 - new_width) / 2
            y = (297 - new_height) / 2

            pdf.image(temp_path, x=x, y=y, w=new_width, h=new_height)
            pdf.output(output_file)

            if os.path.exists(temp_path):
                time.sleep(0.1)
                try:
                    os.remove(temp_path)
                except:
                    pass

            return output_file

        except PermissionError:
            if attempt < max_retries - 1:
                time.sleep(0.5)
                continue
            else:
                raise
        except Exception as e:
            if attempt < max_retries - 1:
                time.sleep(0.5)
                continue
            else:
                raise e


def _dataframe_to_pdf(df: pd.DataFrame, output_file: str, filename: str) -> str:
    """DataFrame ni PDF ga aylantirish"""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Sarlavha
    pdf.set_font("Arial", 'B', 14)
    # Fayl nomini tozalash
    clean_filename = filename.replace('saved_', '').split('_', 1)[-1] if '_' in filename else filename
    pdf.cell(0, 10, f"Ma'lumotlar: {clean_filename}", ln=True, align='C')
    pdf.ln(5)

    # Statistika
    pdf.set_font("Arial", size=9)
    pdf.cell(0, 8, f"Qatorlar: {len(df)} | Ustunlar: {len(df.columns)}", ln=True)
    pdf.ln(3)

    num_cols = len(df.columns)
    available_width = 190
    col_width = min(40, available_width / num_cols)

    # Vertikal format (8+ ustunlar)
    if num_cols > 8:
        pdf.set_font("Arial", size=8)
        for idx, row in df.head(50).iterrows():
            pdf.set_font("Arial", 'B', 9)
            pdf.cell(0, 7, f"Qator {idx + 1}:", ln=True)
            pdf.set_font("Arial", size=8)

            for col in df.columns:
                value = str(row[col])[:50] if pd.notna(row[col]) else '-'
                try:
                    pdf.cell(50, 6, str(col)[:20] + ':', border=0)
                    pdf.cell(0, 6, value, ln=True, border=0)
                except:
                    pdf.cell(50, 6, str(col)[:20] + ':', border=0)
                    pdf.cell(0, 6, '', ln=True, border=0)

            pdf.ln(2)

            if pdf.get_y() > 270:
                pdf.add_page()

        pdf.output(output_file)
        return output_file

    # Jadval format (≤8 ustunlar)
    pdf.set_font("Arial", 'B', 8)
    for col in df.columns:
        col_text = str(col)[:20]
        try:
            pdf.cell(col_width, 8, col_text, border=1, align='C')
        except:
            pdf.cell(col_width, 8, '', border=1, align='C')
    pdf.ln()

    pdf.set_font("Arial", size=7)
    max_rows = min(100, len(df))

    for idx, row in df.head(max_rows).iterrows():
        for item in row:
            text = str(item)[:20] if pd.notna(item) else ''
            try:
                pdf.cell(col_width, 7, text, border=1, align='C')
            except:
                try:
                    clean_text = str(item).encode('latin-1', 'ignore').decode('latin-1')[:20]
                    pdf.cell(col_width, 7, clean_text, border=1, align='C')
                except:
                    pdf.cell(col_width, 7, '', border=1, align='C')
        pdf.ln()

        if pdf.get_y() > 270:
            pdf.add_page()
            pdf.set_font("Arial", 'B', 8)
            for col in df.columns:
                col_text = str(col)[:20]
                try:
                    pdf.cell(col_width, 8, col_text, border=1, align='C')
                except:
                    pdf.cell(col_width, 8, '', border=1, align='C')
            pdf.ln()
            pdf.set_font("Arial", size=7)

    if len(df) > max_rows:
        pdf.ln(5)
        pdf.set_font("Arial", 'I', 9)
        pdf.cell(0, 8, f"... va yana {len(df) - max_rows} qator", ln=True, align='C')

    pdf.output(output_file)
    return output_file


# ==================== Boshqa formatlarga konvertatsiya ====================

def to_csv(filename: str) -> str:
    """Excel/JSON ni CSV ga aylantirish"""
    converter = UniversalConverter()
    file_path = converter.get_file_path(f"saved_{filename}")
    output_file = converter.create_output_filename(file_path, "csv")

    file_ext = filename.split('.')[-1].lower()

    max_retries = 3
    for attempt in range(max_retries):
        try:
            if file_ext in ['xlsx', 'xls']:
                df = converter.safe_read_excel(file_path)
            elif file_ext == 'json':
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                df = pd.DataFrame(data) if isinstance(data, list) else pd.json_normalize(data)
            else:
                raise ValueError(f"CSV ga aylantirish uchun noto'g'ri format: {file_ext}")

            time.sleep(0.2)
            df.to_csv(output_file, index=False, encoding='utf-8')
            return output_file

        except PermissionError:
            if attempt < max_retries - 1:
                time.sleep(0.5)
                continue
            else:
                raise
        except Exception as e:
            if attempt < max_retries - 1:
                time.sleep(0.5)
                continue
            else:
                raise e


def to_excel(filename: str) -> str:
    """CSV/JSON ni Excel ga aylantirish"""
    converter = UniversalConverter()
    file_path = converter.get_file_path(f"saved_{filename}")
    output_file = converter.create_output_filename(file_path, "xlsx")

    file_ext = filename.split('.')[-1].lower()

    max_retries = 3
    for attempt in range(max_retries):
        try:
            if file_ext == 'csv':
                df = pd.read_csv(file_path)
            elif file_ext == 'json':
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                df = pd.DataFrame(data) if isinstance(data, list) else pd.json_normalize(data)
            else:
                raise ValueError(f"Excel ga aylantirish uchun noto'g'ri format: {file_ext}")

            time.sleep(0.2)
            df.to_excel(output_file, index=False, engine='openpyxl')
            return output_file

        except PermissionError:
            if attempt < max_retries - 1:
                time.sleep(0.5)
                continue
            else:
                raise
        except Exception as e:
            if attempt < max_retries - 1:
                time.sleep(0.5)
                continue
            else:
                raise e


def to_json(filename: str) -> str:
    """CSV/Excel ni JSON ga aylantirish"""
    converter = UniversalConverter()
    file_path = converter.get_file_path(f"saved_{filename}")
    output_file = converter.create_output_filename(file_path, "json")

    file_ext = filename.split('.')[-1].lower()

    max_retries = 3
    for attempt in range(max_retries):
        try:
            if file_ext == 'csv':
                df = pd.read_csv(file_path)
            elif file_ext in ['xlsx', 'xls']:
                df = converter.safe_read_excel(file_path)
            else:
                raise ValueError(f"JSON ga aylantirish uchun noto'g'ri format: {file_ext}")

            time.sleep(0.2)
            df.to_json(output_file, orient='records', indent=2, force_ascii=False)
            return output_file

        except PermissionError:
            if attempt < max_retries - 1:
                time.sleep(0.5)
                continue
            else:
                raise
        except Exception as e:
            if attempt < max_retries - 1:
                time.sleep(0.5)
                continue
            else:
                raise e


def to_txt(filename: str) -> str:
    """PDF/DOCX ni TXT ga aylantirish"""
    converter = UniversalConverter()
    file_path = converter.get_file_path(f"saved_{filename}")
    output_file = converter.create_output_filename(file_path, "txt")

    file_ext = filename.split('.')[-1].lower()

    max_retries = 3
    for attempt in range(max_retries):
        try:
            if file_ext == 'pdf':
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text()
            elif file_ext in ['docx', 'doc']:
                doc = docx.Document(file_path)
                text = '\n'.join([para.text for para in doc.paragraphs])
            else:
                raise ValueError(f"TXT ga aylantirish uchun noto'g'ri format: {file_ext}")

            time.sleep(0.2)
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(text)

            return output_file

        except PermissionError:
            if attempt < max_retries - 1:
                time.sleep(0.5)
                continue
            else:
                raise
        except Exception as e:
            if attempt < max_retries - 1:
                time.sleep(0.5)
                continue
            else:
                raise e


Excel_to_pdf = _simple_excel_to_pdf
CSV_to_pdf = CSV_to_pdf
TXT_to_pdf = TXT_to_pdf
Word_to_pdf = _simple_word_to_pdf
